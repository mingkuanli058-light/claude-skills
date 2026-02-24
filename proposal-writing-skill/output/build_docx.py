"""
将建设方案.md 转换为政府信息化方案标准排版的 Word 文档，并自动插入全部图件。

排版规范：
- 页面：A4，上3.7cm 下3.5cm 左2.8cm 右2.6cm
- 正文：方正仿宋_GBK 四号（14pt），行距28磅，首行缩进2字符
- 一级标题（#）：方正小标宋简体 二号（22pt），居中
- 二级标题（##）：方正黑体_GBK 三号（16pt），左对齐
- 三级标题（###）：方正楷体_GBK 四号（14pt），加粗，左对齐
- 四级标题（####）：方正黑体_GBK 四号（14pt），左对齐
- 表格标题行：方正黑体_GBK 小四（12pt），浅蓝底色
- 图件标注：方正黑体_GBK 小四（12pt），居中
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path(__file__).parent
MD_FILE = BASE / "建设方案.md"
OUTPUT = BASE / "执法音视频集中存储与管理系统建设方案.docx"

# 全部图件（按文件名匹配正文中的引用）
CHARTS = {
    "系统总体架构图": BASE / "chart" / "系统总体架构图.png",
    "网络拓扑图":     BASE / "chart" / "网络拓扑图.png",
    "数据流转图":     BASE / "chart" / "数据流转图.png",
    "部署结构图":     BASE / "chart" / "部署结构图.png",
}

# ── Font helpers ──────────────────────────────────────────────────────────

def set_run_font(run, name_cn="方正仿宋_GBK", name_en="Times New Roman", size=None, bold=False, color=None):
    """Set font for a run, with CN and EN font names."""
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def set_paragraph_fmt(p, space_before=0, space_after=0, line_spacing=None, first_line_indent=None, alignment=None):
    """Set paragraph formatting."""
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if line_spacing:
        fmt.line_spacing = Pt(line_spacing)
    if first_line_indent:
        fmt.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        fmt.alignment = alignment

# ── Document setup ────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Page setup: A4, margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)  # 四号
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    style.paragraph_format.line_spacing = Pt(28)  # 28磅行距

    return doc

# ── Cover page ────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Top spacing
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_fmt(p, line_spacing=28)

    # Title line 1
    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=40)
    run = p.add_run("执法音视频集中存储与管理系统")
    set_run_font(run, name_cn="方正小标宋简体", size=26, bold=True)

    # Title line 2
    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_before=0, space_after=24, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=40)
    run = p.add_run("建  设  方  案")
    set_run_font(run, name_cn="方正小标宋简体", size=26, bold=True)

    # Spacing
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_fmt(p, line_spacing=28)

    # Metadata lines (centered)
    meta_lines = [
        ("文件状态：", "交付稿"),
        ("密级等级：", "内部"),
        ("版 本 号：", "V1.0"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        set_paragraph_fmt(p, space_before=6, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=28)
        run = p.add_run(label)
        set_run_font(run, name_cn="方正黑体_GBK", size=14)
        run = p.add_run(value)
        set_run_font(run, name_cn="方正仿宋_GBK", size=14)

    doc.add_page_break()

# ── Parse markdown ────────────────────────────────────────────────────────

def parse_md(filepath):
    """Parse markdown into a list of elements:
    ('h1', text), ('h2', text), ('h3', text),
    ('para', text), ('bullet', text), ('hr',),
    ('table', headers, rows)
    """
    text = filepath.read_text(encoding='utf-8')
    lines = text.split('\n')
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Headers (check h4 before h3)
        if line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
            i += 1
            continue
        if line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        if line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2  # skip separator
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            elements.append(('table', headers, rows))
            continue

        # Bullet
        if line.startswith('- ') or line.startswith('  - '):
            indent_level = 0
            if line.startswith('  - '):
                indent_level = 1
            elements.append(('bullet', line.lstrip(' -').strip(), indent_level))
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            elements.append(('numbered', m.group(2).strip(), int(m.group(1))))
            i += 1
            continue

        # Paragraph (non-empty)
        if line.strip():
            elements.append(('para', line.strip()))
            i += 1
            continue

        i += 1

    return elements

# ── Render helpers ────────────────────────────────────────────────────────

def render_inline(paragraph, text, default_cn="方正仿宋_GBK", default_size=14, default_bold=False):
    """Render text with **bold** inline markup."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_run_font(run, name_cn="方正黑体_GBK", size=default_size, bold=True)
        else:
            # Handle `code` backticks too
            subparts = re.split(r'(`.*?`)', part)
            for sp in subparts:
                if sp.startswith('`') and sp.endswith('`'):
                    run = paragraph.add_run(sp[1:-1])
                    set_run_font(run, name_cn="方正仿宋_GBK", name_en="Consolas", size=default_size)
                else:
                    if sp:
                        run = paragraph.add_run(sp)
                        set_run_font(run, name_cn=default_cn, size=default_size, bold=default_bold)

def add_table(doc, headers, rows):
    """Add a formatted table."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=24)
        run = p.add_run(h)
        set_run_font(run, name_cn="方正黑体_GBK", size=12, bold=True)
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j >= ncols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=24)
            render_inline(p, val, default_size=12)

    # Add spacing after table
    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_before=6, line_spacing=10)

def insert_chart(doc, chart_name, image_path):
    """Insert a chart image with caption."""
    # Caption above
    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_before=12, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=28)
    run = p.add_run(f"图：{chart_name}")
    set_run_font(run, name_cn="方正黑体_GBK", size=12, bold=True)

    # Image - fit to page width (approx 15.6 cm usable)
    p = doc.add_paragraph()
    set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.6))

    # Spacing after
    p = doc.add_paragraph()
    set_paragraph_fmt(p, line_spacing=10)

# ── Main build ────────────────────────────────────────────────────────────

def build():
    doc = setup_document()
    add_cover_page(doc)

    elements = parse_md(MD_FILE)

    for elem in elements:
        etype = elem[0]

        if etype == 'h1':
            # 文档总标题：方正小标宋简体 二号 居中
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=24, space_after=18,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=36)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正小标宋简体", size=22, bold=True)

        elif etype == 'h2':
            # 章标题（一、二、…）：方正黑体 三号 左对齐
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=18, space_after=12,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=32)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正黑体_GBK", size=16, bold=True)

        elif etype == 'h3':
            # 节标题（（一）（二）…）：方正楷体 四号 加粗 左对齐
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=12, space_after=6,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=28)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正楷体_GBK", size=14, bold=True)

        elif etype == 'h4':
            # 子节标题（1、2、…）：方正黑体 四号 左对齐
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=8, space_after=4,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=28)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正黑体_GBK", size=14, bold=False)

        elif etype == 'para':
            text = elem[1]

            # 检测图件引用行，替换为实际图片插入
            chart_inserted = False
            for chart_name, chart_path in CHARTS.items():
                if chart_name in text and '详见' in text:
                    insert_chart(doc, chart_name, chart_path)
                    chart_inserted = True
                    break
            if chart_inserted:
                continue

            p = doc.add_paragraph()
            set_paragraph_fmt(p, first_line_indent=0.74, line_spacing=28)
            render_inline(p, text)

        elif etype == 'bullet':
            text = elem[1]
            indent_level = elem[2] if len(elem) > 2 else 0
            p = doc.add_paragraph()
            left_indent = 0.74 + indent_level * 0.74
            set_paragraph_fmt(p, line_spacing=28)
            p.paragraph_format.left_indent = Cm(left_indent)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run("— ")
            set_run_font(run, size=14)
            render_inline(p, text)

        elif etype == 'numbered':
            text = elem[1]
            num = elem[2]
            p = doc.add_paragraph()
            set_paragraph_fmt(p, line_spacing=28)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(f"{num}. ")
            set_run_font(run, size=14)
            render_inline(p, text)

        elif etype == 'table':
            add_table(doc, elem[1], elem[2])

    doc.save(str(OUTPUT))
    print(f"Done: {OUTPUT}")

if __name__ == '__main__':
    build()
