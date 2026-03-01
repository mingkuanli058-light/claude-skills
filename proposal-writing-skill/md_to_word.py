#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为符合政府规范的Word文档
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_row_height(row, height):
    """设置行高"""
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:type'), 'auto')
    trPr.append(trHeight)

def create_government_word_doc(md_file, output_file, chart_dir):
    """
    创建符合政府规范的Word文档
    """
    # 创建Document对象
    doc = Document()

    # 设置文档宽度、页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 分割内容为行
    lines = content.split('\n')

    i = 0
    current_list_type = None  # 追踪当前列表类型

    while i < len(lines):
        line = lines[i]

        # 处理标题
        if line.startswith('# '):
            # 一级标题
            title = line.replace('# ', '').strip()
            p = doc.add_paragraph(title, style='Heading 1')
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            p_format.line_spacing = 1.5

            # 设置标题格式
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.name = '仿宋'

        elif line.startswith('## '):
            # 二级标题
            title = line.replace('## ', '').strip()
            p = doc.add_paragraph(title, style='Heading 2')
            p_format = p.paragraph_format
            p_format.space_before = Pt(10)
            p_format.space_after = Pt(8)
            p_format.left_indent = Cm(0.5)

            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.name = '仿宋'

        elif line.startswith('### '):
            # 三级标题
            title = line.replace('### ', '').strip()
            p = doc.add_paragraph(title)
            p_format = p.paragraph_format
            p_format.left_indent = Cm(1.0)
            p_format.space_before = Pt(8)
            p_format.space_after = Pt(6)

            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.name = '仿宋'

        elif line.startswith('#### '):
            # 四级标题
            title = line.replace('#### ', '').strip()
            p = doc.add_paragraph(title)
            p_format = p.paragraph_format
            p_format.left_indent = Cm(1.5)

            for run in p.runs:
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = '仿宋'

        # 处理表格
        elif line.strip().startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退一行

            if len(table_lines) >= 2:
                # 解析表格
                headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]

                # 创建表格
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'

                # 设置表头
                header_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    # 设置表头格式
                    set_cell_background(header_cells[idx], 'D3D3D3')

                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.bold = True
                            run.font.name = '仿宋'
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 添加表格行（跳过分隔符行）
                for line_idx in range(2, len(table_lines)):
                    row_data = [cell.strip() for cell in table_lines[line_idx].split('|')[1:-1]]

                    # 跳过分隔符行
                    if all(c in '-:|' for c in table_lines[line_idx].replace(' ', '')):
                        continue

                    row_cells = table.add_row().cells
                    for idx, cell_data in enumerate(row_data):
                        row_cells[idx].text = cell_data
                        for paragraph in row_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = '仿宋'
                            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 处理分隔线
        elif line.strip().startswith('---'):
            doc.add_paragraph()  # 添加空行

        # 处理列表项
        elif line.strip().startswith('- '):
            list_item = line.replace('- ', '').strip()
            p = doc.add_paragraph(list_item, style='List Bullet')
            p_format = p.paragraph_format
            p_format.left_indent = Cm(0.75)
            p_format.first_line_indent = Cm(-0.25)
            p_format.space_after = Pt(3)

            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '仿宋'
            current_list_type = 'bullet'

        # 处理有序列表
        elif line.strip() and line.strip()[0].isdigit() and '.' in line.strip()[:3]:
            list_item = re.sub(r'^\d+\.\s*', '', line).strip()
            p = doc.add_paragraph(list_item, style='List Number')
            p_format = p.paragraph_format
            p_format.left_indent = Cm(0.75)
            p_format.space_after = Pt(3)

            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '仿宋'
            current_list_type = 'number'

        # 处理图表引用
        elif '详见' in line and ('output/chart/' in line or 'chart/' in line):
            # 提取图表路径
            match = re.search(r'`([^`]*chart[^`]*)`', line)
            if match:
                chart_path = match.group(1)
                chart_name = chart_path.split('/')[-1]

                # 尝试找到对应的PNG文件
                png_file = None
                if chart_name.endswith('.html'):
                    png_name = chart_name.replace('.html', '.png')
                else:
                    png_name = chart_name

                png_path = os.path.join(chart_dir, png_name)

                if os.path.exists(png_path):
                    # 添加说明文本
                    p = doc.add_paragraph(line.replace(f'`{chart_path}`', f'（详见下图）'))
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = '仿宋'

                    # 添加图表
                    doc.add_picture(png_path, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # 如果PNG不存在，只添加文本
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = '仿宋'
            else:
                # 普通文本行
                if line.strip():
                    p = doc.add_paragraph(line)
                    p_format = p.paragraph_format
                    p_format.first_line_indent = Cm(0.5)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = '仿宋'

        # 普通文本行
        elif line.strip() and not line.strip().startswith('**'):
            p = doc.add_paragraph(line)
            p_format = p.paragraph_format
            if not current_list_type:
                p_format.first_line_indent = Cm(0.5)

            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '仿宋'

        # 处理加粗文本行
        elif line.strip().startswith('**'):
            line_text = line.replace('**', '')
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.first_line_indent = Cm(0.5)

            run = p.add_run(line_text)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = '仿宋'

        i += 1

    # 保存文档
    doc.save(output_file)
    print(f"Word文档已生成：{output_file}")

def main():
    # 文件路径
    md_file = r'd:\work\Claude-Skill\proposal-writing-skill\output\建设方案.md'
    output_file = r'd:\work\Claude-Skill\proposal-writing-skill\output\建设方案.docx'
    chart_dir = r'd:\work\Claude-Skill\proposal-writing-skill\output\chart'

    # 检查输入文件
    if not os.path.exists(md_file):
        print(f"错误：找不到Markdown文件 {md_file}")
        return

    if not os.path.exists(chart_dir):
        print(f"错误：找不到图表目录 {chart_dir}")
        return

    # 生成Word文档
    create_government_word_doc(md_file, output_file, chart_dir)

if __name__ == '__main__':
    main()
