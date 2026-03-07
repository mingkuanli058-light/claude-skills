"""
将 output/proposal.md 转换为政府标准排版 Word 文档。

版式规范（来自用户要求）：
  纸张：A4，页边距 2.5cm（上下左右）
  正文：仿宋_GB2312，三号（16pt），固定行距28磅，首行缩进2字符，两端对齐
  一级标题（##）：黑体，三号（16pt）
  二级标题（###）：黑体，四号（14pt）
  三级标题（####）：宋体，小四（12pt）
  目录：显示至三级标题，页码右对齐
  图件：居中，图下添加"图X-X 标题"，保持原比例
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

BASE   = Path(__file__).parent
MD_IN  = BASE / "proposal.md"
DOCX_OUT = BASE / "建设方案.docx"
CHART_DIR = BASE / "chart"

# 图件：blockquote ID → (PNG文件名, 显示名称)
CHART_MAP = {
    "A-01": ("system_architecture.png",  "系统总体架构图"),
    "A-02": ("network_topology.png",      "网络拓扑图"),
    "A-03": ("data_flow.png",             "数据流转图"),
    "A-04": ("deployment_structure.png",  "部署结构图"),
    "A-05": ("server_role_structure.png", "服务器角色划分图"),
    "B-01": ("project_gantt.png",         "实施甘特图"),
}

# 章节编号 → 图件计数器（用于 "图5-1" 格式）
_fig_counters = {}
_current_chapter = 0
_fig_index = 0


# ─────────────────────────────────────────────────────────
# 字体辅助
# ─────────────────────────────────────────────────────────

def _set_font(run, cn_name, en_name="Times New Roman", size_pt=None, bold=False, italic=False, color=None):
    """设置 run 的中西文字体、字号、加粗、斜体。"""
    run.bold = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = en_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'),    en_name)
    rFonts.set(qn('w:hAnsi'),    en_name)


def _set_para_fmt(p,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_exact_pt=28,
                  first_indent_pt=None,
                  left_indent_cm=None,
                  space_before_pt=0,
                  space_after_pt=0):
    """统一设置段落格式（固定行距、缩进、对齐）。"""
    fmt = p.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after  = Pt(space_after_pt)
    if line_exact_pt is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_exact_pt)
    if first_indent_pt is not None:
        fmt.first_line_indent = Pt(first_indent_pt)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)


# ─────────────────────────────────────────────────────────
# 文档初始化
# ─────────────────────────────────────────────────────────

def _setup_doc():
    doc = Document()
    # A4 + 2.5cm 四边
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    # 修改 Normal 样式默认字体（仿宋_GB2312，三号=16pt）
    normal = doc.styles['Normal']
    normal.font.size = Pt(16)
    normal.font.name = 'Times New Roman'
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    # 修改 Heading 1/2/3 样式，使 TOC 能正常识别
    for style_name, cn_font, size_pt in [
        ('Heading 1', '黑体', 16),
        ('Heading 2', '黑体', 14),
        ('Heading 3', '宋体', 12),
    ]:
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.size = Pt(size_pt)
        st.font.name = 'Times New Roman'
        st.font.bold = True if style_name in ('Heading 1', 'Heading 2') else False
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after  = Pt(0)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        st.paragraph_format.line_spacing = Pt(28)
        rPr2 = st.element.get_or_add_rPr()
        rFonts2 = rPr2.get_or_add_rFonts()
        rFonts2.set(qn('w:eastAsia'), cn_font)

    return doc


# ─────────────────────────────────────────────────────────
# 目录页
# ─────────────────────────────────────────────────────────

def _add_toc(doc):
    r"""插入自动目录（\o 1-3 = 显示至三级，\z = 隐藏页码地址，\u = 用段落大纲级别）。"""
    # 目录标题
    p = doc.add_paragraph()
    _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line_exact_pt=28, space_before_pt=0, space_after_pt=12)
    run = p.add_run("目  录")
    _set_font(run, cn_name='黑体', size_pt=16, bold=True)

    # TOC 域代码
    p2 = doc.add_paragraph()
    _set_para_fmt(p2, line_exact_pt=28)
    # 插入 BEGIN 域
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    r1 = OxmlElement('w:r')
    r1.append(fldBegin)
    p2._p.append(r1)

    # 域指令
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r' TOC \o "1-3" \h \z \u '
    r2 = OxmlElement('w:r')
    r2.append(instrText)
    p2._p.append(r2)

    # SEPARATE
    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    r3 = OxmlElement('w:r')
    r3.append(fldSep)
    p2._p.append(r3)

    # 占位文字
    t = OxmlElement('w:t')
    t.text = '（请在 Word 中右键 → 更新域以生成目录）'
    r4 = OxmlElement('w:r')
    r4.append(t)
    p2._p.append(r4)

    # END
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    r5 = OxmlElement('w:r')
    r5.append(fldEnd)
    p2._p.append(r5)

    # 目录后分页
    doc.add_page_break()


# ─────────────────────────────────────────────────────────
# 内联富文本渲染（**bold**、`code`、*italic*）
# ─────────────────────────────────────────────────────────

def _render_inline(p, text, cn_font='仿宋_GB2312', size_pt=16):
    """解析 **粗体**、`代码`、*斜体* 并逐 run 输出到段落。"""
    # 去除首尾 Markdown 图片/链接（blockquote后续行已被过滤）
    text = text.strip()
    # 先按 **...** 切分
    segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for seg in segments:
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            run = p.add_run(seg[2:-2])
            _set_font(run, cn_name='黑体', size_pt=size_pt, bold=True)
        elif seg.startswith('*') and seg.endswith('*'):
            run = p.add_run(seg[1:-1])
            _set_font(run, cn_name='仿宋_GB2312', size_pt=size_pt, italic=True)
        elif seg.startswith('`') and seg.endswith('`'):
            run = p.add_run(seg[1:-1])
            _set_font(run, cn_name='仿宋_GB2312', en_name='Consolas', size_pt=size_pt-2)
        else:
            run = p.add_run(seg)
            _set_font(run, cn_name=cn_font, size_pt=size_pt)


# ─────────────────────────────────────────────────────────
# 图件插入
# ─────────────────────────────────────────────────────────

def _insert_chart(doc, chart_id, chart_name, fig_label):
    """插入图件 PNG，图下添加编号标题。"""
    png_name, _ = CHART_MAP[chart_id]
    png_path = CHART_DIR / png_name
    if not png_path.exists():
        # 占位段落
        p = doc.add_paragraph()
        _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_exact_pt=28)
        run = p.add_run(f"【图件未找到：{png_name}】")
        _set_font(run, cn_name='仿宋_GB2312', size_pt=12, color=(180, 0, 0))
    else:
        # 图片段落：居中，宽度 = 版心（21 - 2.5×2 = 16cm）
        p = doc.add_paragraph()
        _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line_exact_pt=28, space_before_pt=6, space_after_pt=0)
        run = p.add_run()
        run.add_picture(str(png_path), width=Cm(16))

    # 图下标题：图X-X 名称
    caption_p = doc.add_paragraph()
    _set_para_fmt(caption_p, align=WD_ALIGN_PARAGRAPH.CENTER,
                  line_exact_pt=28, space_before_pt=4, space_after_pt=8)
    run = caption_p.add_run(f"{fig_label}　{chart_name}")
    _set_font(run, cn_name='黑体', size_pt=12, bold=True)


# ─────────────────────────────────────────────────────────
# 表格
# ─────────────────────────────────────────────────────────

def _add_table(doc, headers, rows):
    ncols = max(len(headers), 1)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头行
    for j, h in enumerate(headers):
        if j >= ncols:
            break
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_exact_pt=22)
        run = p.add_run(h)
        _set_font(run, cn_name='黑体', size_pt=11, bold=True)
        # 浅蓝表头底色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:color="auto" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j >= ncols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_exact_pt=22)
            _render_inline(p, val, cn_font='仿宋_GB2312', size_pt=11)

    # 表后间距
    sp = doc.add_paragraph()
    _set_para_fmt(sp, line_exact_pt=14, space_before_pt=0, space_after_pt=0)


# ─────────────────────────────────────────────────────────
# Markdown 解析
# ─────────────────────────────────────────────────────────

def _parse_md(filepath):
    """
    解析 Markdown → 元素列表。元素类型：
      ('h1', text)
      ('h2', text)
      ('h3', text)
      ('h4', text)
      ('para', text)
      ('bullet', text, indent_level)
      ('numbered', num_str, text)
      ('table', headers, rows)
      ('chart', chart_id, chart_name)   ← blockquote 图件引用
      ('blockquote_skip',)              ← 图件后续说明行，跳过
      ('hr',)
    """
    text = filepath.read_text(encoding='utf-8')
    lines = text.splitlines()
    elements = []
    i = 0
    in_blockquote_chart = False  # 是否正在处理图件 blockquote 的后续行

    while i < len(lines):
        line = lines[i]

        # --- 水平线 ---
        if re.match(r'^-{3,}$', line.strip()):
            elements.append(('hr',))
            i += 1
            in_blockquote_chart = False
            continue

        # --- blockquote ---
        if line.startswith('>'):
            content = line[1:].strip()
            # 图件引用行格式：**图件 X-XX**：名称
            m = re.match(r'\*\*图件\s+([A-Z]-\d+)\*\*[：:]\s*(.+)', content)
            if m:
                chart_id   = m.group(1).strip()
                chart_name = m.group(2).strip()
                elements.append(('chart', chart_id, chart_name))
                in_blockquote_chart = True
            else:
                # 图件后续说明行或普通 blockquote → 跳过
                if in_blockquote_chart:
                    elements.append(('blockquote_skip',))
                else:
                    # 普通 blockquote 作为正文段落
                    elements.append(('para', content))
            i += 1
            continue

        # 非 blockquote 行，重置标志
        in_blockquote_chart = False

        # --- 标题 ---
        if line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
            i += 1
            continue
        if line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        if line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue

        # --- 表格 ---
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|?\s*[-:]+', lines[i + 1]):
            headers = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2  # 跳过分隔行
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            elements.append(('table', headers, rows))
            continue

        # --- 列表项 ---
        m_bullet = re.match(r'^(\s*)[-*]\s+(.+)', line)
        if m_bullet:
            indent = len(m_bullet.group(1)) // 2
            elements.append(('bullet', m_bullet.group(2).strip(), indent))
            i += 1
            continue

        # --- 有序列表 ---
        m_num = re.match(r'^(\d+)[.)]\s+(.+)', line)
        if m_num:
            elements.append(('numbered', m_num.group(1), m_num.group(2).strip()))
            i += 1
            continue

        # --- 普通段落 ---
        if line.strip():
            elements.append(('para', line.strip()))
            i += 1
            continue

        i += 1

    return elements


# ─────────────────────────────────────────────────────────
# 图件编号辅助（图5-1 格式）
# ─────────────────────────────────────────────────────────

_chapter_num = 0       # 当前一级章节编号（##）
_fig_counter  = {}     # {chapter_num: count}

def _next_fig_label(chapter_num):
    """返回下一个图件编号，如 "图5-1"。"""
    _fig_counter.setdefault(chapter_num, 0)
    _fig_counter[chapter_num] += 1
    return f"图{chapter_num}-{_fig_counter[chapter_num]}"


# ─────────────────────────────────────────────────────────
# 主构建流程
# ─────────────────────────────────────────────────────────

def build():
    global _chapter_num

    doc = _setup_doc()
    _add_toc(doc)

    elements = _parse_md(MD_IN)

    for elem in elements:
        etype = elem[0]

        # ── 标题：# (文档总标题)
        if etype == 'h1':
            p = doc.add_paragraph(style='Heading 1')
            _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                          line_exact_pt=36, space_before_pt=12, space_after_pt=18)
            run = p.add_run(elem[1])
            _set_font(run, cn_name='黑体', size_pt=22, bold=True)

        # ── 一级标题：## → Heading 1（黑体，三号16pt）
        elif etype == 'h2':
            _chapter_num += 1
            p = doc.add_paragraph(style='Heading 1')
            _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                          line_exact_pt=28, space_before_pt=18, space_after_pt=10)
            run = p.add_run(elem[1])
            _set_font(run, cn_name='黑体', size_pt=16, bold=True)

        # ── 二级标题：### → Heading 2（黑体，四号14pt）
        elif etype == 'h3':
            p = doc.add_paragraph(style='Heading 2')
            _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                          line_exact_pt=28, space_before_pt=12, space_after_pt=6)
            run = p.add_run(elem[1])
            _set_font(run, cn_name='黑体', size_pt=14, bold=True)

        # ── 三级标题：#### → Heading 3（宋体，小四12pt）
        elif etype == 'h4':
            p = doc.add_paragraph(style='Heading 3')
            _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                          line_exact_pt=28, space_before_pt=8, space_after_pt=4)
            run = p.add_run(elem[1])
            _set_font(run, cn_name='宋体', size_pt=12, bold=False)

        # ── 正文段落（仿宋_GB2312，三号16pt，首行缩进2字符=32pt，两端对齐）
        elif etype == 'para':
            p = doc.add_paragraph()
            _set_para_fmt(p,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_exact_pt=28,
                          first_indent_pt=32,   # 2字符 @ 16pt
                          space_before_pt=0,
                          space_after_pt=0)
            _render_inline(p, elem[1], cn_font='仿宋_GB2312', size_pt=16)

        # ── 无序列表
        elif etype == 'bullet':
            text        = elem[1]
            indent_lvl  = elem[2] if len(elem) > 2 else 0
            p = doc.add_paragraph()
            left_cm = 0.8 + indent_lvl * 0.8
            _set_para_fmt(p,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_exact_pt=28,
                          left_indent_cm=left_cm)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run("● ")
            _set_font(run, cn_name='仿宋_GB2312', size_pt=16)
            _render_inline(p, text, cn_font='仿宋_GB2312', size_pt=16)

        # ── 有序列表
        elif etype == 'numbered':
            num  = elem[1]
            text = elem[2]
            p = doc.add_paragraph()
            _set_para_fmt(p,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_exact_pt=28,
                          left_indent_cm=0.8)
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(f"{num}. ")
            _set_font(run, cn_name='仿宋_GB2312', size_pt=16)
            _render_inline(p, text, cn_font='仿宋_GB2312', size_pt=16)

        # ── 表格
        elif etype == 'table':
            _add_table(doc, elem[1], elem[2])

        # ── 图件引用 → 插入 PNG
        elif etype == 'chart':
            chart_id   = elem[1]
            chart_name = elem[2]
            if chart_id in CHART_MAP:
                label = _next_fig_label(_chapter_num)
                _insert_chart(doc, chart_id, chart_name, label)
            else:
                # 未注册的图件 ID，作为普通说明段落
                p = doc.add_paragraph()
                _set_para_fmt(p, line_exact_pt=28, first_indent_pt=32)
                _render_inline(p, f"【图件 {chart_id}：{chart_name}】",
                               cn_font='仿宋_GB2312', size_pt=14)

        # ── 图件后续说明行，跳过
        elif etype == 'blockquote_skip':
            continue

        # ── 水平线：空行隔断
        elif etype == 'hr':
            p = doc.add_paragraph()
            _set_para_fmt(p, line_exact_pt=14)

    doc.save(str(DOCX_OUT))
    print(f"\nDone: {DOCX_OUT}")


if __name__ == '__main__':
    build()
