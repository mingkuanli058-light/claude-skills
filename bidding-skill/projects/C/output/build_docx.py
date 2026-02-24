"""
将投标文件 Markdown 转换为政府采购投标文件标准排版的 Word 文档。

排版规范：
- 页面：A4，上3.7cm 下3.5cm 左2.8cm 右2.6cm
- 正文：方正仿宋_GBK 四号（14pt），行距28磅，首行缩进2字符
- 一级标题（#）：方正小标宋简体 二号（22pt），居中
- 二级标题（##）：方正黑体_GBK 三号（16pt），左对齐
- 三级标题（###）：方正楷体_GBK 四号（14pt），加粗，左对齐
- 四级标题（####）：方正黑体_GBK 四号（14pt），左对齐
- 表格标题行：方正黑体_GBK 小四（12pt），浅蓝底色

用法：
    python build_docx.py 技术方案.md
    python build_docx.py 商务方案.md
    python build_docx.py 报价文件.md
    python build_docx.py 资格文件.md
    python build_docx.py --all  # 构建全部
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path(__file__).parent
CHART_DIR = BASE / "chart"

# 可构建的文件列表
BUILD_TARGETS = [
    "技术方案.md",
    "商务方案.md",
    "报价文件.md",
    "资格文件.md",
]


# ── Font helpers ──

def set_run_font(run, name_cn="方正仿宋_GBK", name_en="Times New Roman", size=None, bold=False, color=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def set_paragraph_fmt(p, space_before=0, space_after=0, line_spacing=None, first_line_indent=None, alignment=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if line_spacing:
        fmt.line_spacing = Pt(line_spacing)
    if first_line_indent:
        fmt.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        fmt.alignment = alignment


# ── Document setup ──

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    style.paragraph_format.line_spacing = Pt(28)

    return doc


# ── Parse markdown ──

def parse_md(filepath):
    text = filepath.read_text(encoding='utf-8')
    lines = text.split('\n')
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == '---':
            i += 1
            continue

        if line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
            i += 1
            continue
        if line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        if line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue

        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            elements.append(('table', headers, rows))
            continue

        if line.startswith('- ') or line.startswith('  - '):
            indent_level = 1 if line.startswith('  - ') else 0
            elements.append(('bullet', line.lstrip(' -').strip(), indent_level))
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            elements.append(('numbered', m.group(2).strip(), int(m.group(1))))
            i += 1
            continue

        if line.strip():
            elements.append(('para', line.strip()))
            i += 1
            continue

        i += 1

    return elements


# ── Render helpers ──

def render_inline(paragraph, text, default_cn="方正仿宋_GBK", default_size=14, default_bold=False):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_run_font(run, name_cn="方正黑体_GBK", size=default_size, bold=True)
        else:
            subparts = re.split(r'(`.*?`)', part)
            for sp in subparts:
                if sp.startswith('`') and sp.endswith('`'):
                    run = paragraph.add_run(sp[1:-1])
                    set_run_font(run, name_cn="方正仿宋_GBK", name_en="Consolas", size=default_size)
                else:
                    if sp:
                        run = paragraph.add_run(sp)
                        set_run_font(run, name_cn=default_cn, size=default_size, bold=default_bold)


def add_table(doc, headers, rows):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=24)
        run = p.add_run(h)
        set_run_font(run, name_cn="方正黑体_GBK", size=12, bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._element.get_or_add_tcPr().append(shading)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            if j >= ncols:
                break
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=24)
            render_inline(p, val, default_size=12)

    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_before=6, line_spacing=10)


def insert_chart(doc, chart_name, image_path):
    p = doc.add_paragraph()
    set_paragraph_fmt(p, space_before=12, space_after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=28)
    run = p.add_run(f"图：{chart_name}")
    set_run_font(run, name_cn="方正黑体_GBK", size=12, bold=True)

    p = doc.add_paragraph()
    set_paragraph_fmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.6))

    p = doc.add_paragraph()
    set_paragraph_fmt(p, line_spacing=10)


# ── Main build ──

def build_one(md_filename):
    md_path = BASE / md_filename
    if not md_path.exists():
        print(f"Skip: {md_path} not found")
        return

    doc = setup_document()
    elements = parse_md(md_path)

    # 收集 chart/ 下所有 PNG
    charts = {}
    if CHART_DIR.exists():
        for png in CHART_DIR.glob("*.png"):
            charts[png.stem] = png

    for elem in elements:
        etype = elem[0]

        if etype == 'h1':
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=24, space_after=18,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=36)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正小标宋简体", size=22, bold=True)

        elif etype == 'h2':
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=18, space_after=12,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=32)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正黑体_GBK", size=16, bold=True)

        elif etype == 'h3':
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=12, space_after=6,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=28)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正楷体_GBK", size=14, bold=True)

        elif etype == 'h4':
            p = doc.add_paragraph()
            set_paragraph_fmt(p, space_before=8, space_after=4,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=28)
            run = p.add_run(elem[1])
            set_run_font(run, name_cn="方正黑体_GBK", size=14, bold=False)

        elif etype == 'para':
            text = elem[1]

            # 检测图件引用行
            chart_inserted = False
            for chart_name, chart_path in charts.items():
                if chart_name in text and '详见' in text:
                    insert_chart(doc, chart_name, chart_path)
                    chart_inserted = True
                    break
            if chart_inserted:
                continue

            p = doc.add_paragraph()
            set_paragraph_fmt(p, first_line_indent=0.74, line_spacing=28)
            render_inline(p, text)

        elif etype == 'bullet':
            text = elem[1]
            indent_level = elem[2] if len(elem) > 2 else 0
            p = doc.add_paragraph()
            left_indent = 0.74 + indent_level * 0.74
            set_paragraph_fmt(p, line_spacing=28)
            p.paragraph_format.left_indent = Cm(left_indent)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run("— ")
            set_run_font(run, size=14)
            render_inline(p, text)

        elif etype == 'numbered':
            text = elem[1]
            num = elem[2]
            p = doc.add_paragraph()
            set_paragraph_fmt(p, line_spacing=28)
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(f"{num}. ")
            set_run_font(run, size=14)
            render_inline(p, text)

        elif etype == 'table':
            add_table(doc, elem[1], elem[2])

    out_name = md_filename.replace('.md', '.docx')
    out_path = BASE / out_name
    doc.save(str(out_path))
    print(f"Done: {out_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python build_docx.py <文件名.md> | --all")
        print(f"可用目标: {', '.join(BUILD_TARGETS)}")
        return

    if sys.argv[1] == '--all':
        for t in BUILD_TARGETS:
            build_one(t)
    else:
        build_one(sys.argv[1])


if __name__ == '__main__':
    main()
